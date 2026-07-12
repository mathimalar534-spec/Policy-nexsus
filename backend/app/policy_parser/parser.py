import os
import re
from typing import Dict, Any
from pypdf import PdfReader
from docx import Document

class PolicyParser:
    @staticmethod
    def normalize_text(text: str) -> str:
        """
        Cleans and normalizes raw text:
        - Replaces multiple spaces with a single space.
        - Normalizes newlines.
        - Removes non-printable or corrupt control characters.
        """
        if not text:
            return ""
        # Remove control characters except newline and tab
        text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E]', '', text)
        # Normalize whitespace (replace multiple spaces with one)
        text = re.sub(r'[ \t]+', ' ', text)
        # Normalize newlines (max 2 consecutive newlines)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    @classmethod
    def parse_pdf(cls, file_path: str) -> str:
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            raise RuntimeError(f"Error parsing PDF file: {str(e)}")
        return cls.normalize_text(text)

    @classmethod
    def parse_docx(cls, file_path: str) -> str:
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise RuntimeError(f"Error parsing DOCX file: {str(e)}")
        return cls.normalize_text(text)

    @classmethod
    def parse_txt(cls, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        except Exception as e:
            raise RuntimeError(f"Error reading TXT file: {str(e)}")
        return cls.normalize_text(text)

    @classmethod
    def parse_markdown(cls, file_path: str) -> str:
        # Standard markdown is parsed similarly to txt, with minor structural cleaning
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            # Remove Markdown syntax characters if we want pure text normalization,
            # but keep it readable. For now, normal TXT cleaning is sufficient and keeps context intact.
        except Exception as e:
            raise RuntimeError(f"Error reading Markdown file: {str(e)}")
        return cls.normalize_text(text)

    @classmethod
    def extract_text(cls, file_path: str, file_type: str) -> str:
        ext = file_type.lower().strip('.')
        if ext == "pdf":
            return cls.parse_pdf(file_path)
        elif ext in ["docx", "doc"]:
            return cls.parse_docx(file_path)
        elif ext == "txt":
            return cls.parse_txt(file_path)
        elif ext in ["md", "markdown"]:
            return cls.parse_markdown(file_path)
        else:
            # Try to parse as txt by default
            return cls.parse_txt(file_path)
